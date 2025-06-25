"""
Text extraction module for various document formats.
"""

import logging
from pathlib import Path
from typing import Optional, Dict, Any
import io


try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("PDF libraries not available. Install PyPDF2 and pdfplumber for PDF support.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install python-docx for DOCX support.")

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("OCR libraries not available. Install Pillow and pytesseract for OCR support.")

from utils import clean_text, get_file_info
from config import TESSERACT_CMD, OCR_LANGUAGES

class TextExtractor:
    """
    Text extraction class for various document formats.
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
       
        if OCR_AVAILABLE:
            try:
                pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
            except Exception as e:
                self.logger.warning(f"Could not configure tesseract: {str(e)}")
    
    def extract_text(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from various file formats.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        result = {
            'success': False,
            'text': '',
            'page_count': 0,
            'word_count': 0,
            'character_count': 0,
            'extraction_method': '',
            'error': None
        }
        
        try:
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.txt':
                result = self._extract_from_txt(file_path)
            elif file_extension == '.pdf':
                result = self._extract_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                result = self._extract_from_docx(file_path)
            else:
                result['error'] = f"Unsupported file format: {file_extension}"
                self.logger.error(result['error'])
                return result
            
        
            if result['text']:
                result['character_count'] = len(result['text'])
                result['word_count'] = len(result['text'].split())
            
            result['success'] = True
            self.logger.info(f"Successfully extracted text from {file_path}")
            
        except Exception as e:
            error_msg = f"Error extracting text from {file_path}: {str(e)}"
            result['error'] = error_msg
            self.logger.error(error_msg)
        
        return result
    
    def _extract_from_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT files."""
        result = {'extraction_method': 'direct_text_read', 'page_count': 1}
        
       
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                result['text'] = clean_text(text)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode text file with any supported encoding")
        
        return result
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF files."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF libraries not available")
        
        result = {'extraction_method': 'pdf_text_extraction'}
        text_content = []
        

        try:
            with pdfplumber.open(file_path) as pdf:
                result['page_count'] = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                
                if text_content:
                    result['text'] = clean_text('\n'.join(text_content))
                    result['extraction_method'] = 'pdfplumber'
                    return result
        except Exception as e:
            self.logger.warning(f"pdfplumber failed: {str(e)}, trying PyPDF2")

        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                result['page_count'] = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                
                if text_content:
                    result['text'] = clean_text('\n'.join(text_content))
                    result['extraction_method'] = 'pypdf2'
                    return result
        except Exception as e:
            self.logger.warning(f"PyPDF2 failed: {str(e)}")
        
        if not text_content and OCR_AVAILABLE:
            self.logger.info("No text extracted from PDF, attempting OCR")
            result = self._extract_with_ocr(file_path)
            result['extraction_method'] = 'pdf_ocr'
        else:
            result['text'] = ''
        
        return result
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available")
        
        result = {'extraction_method': 'docx_text_extraction', 'page_count': 1}
        
        doc = Document(file_path)
        text_content = []
        

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        result['text'] = clean_text('\n'.join(text_content))
        return result
    
    def _extract_with_ocr(self, file_path: Path) -> Dict[str, Any]:
        """Extract text using OCR."""
        if not OCR_AVAILABLE:
            raise ImportError("OCR libraries not available")
        
        result = {'extraction_method': 'ocr', 'page_count': 0}
        
        try:
            if file_path.suffix.lower() == '.pdf':
              
                result['text'] = "OCR extraction requires pdf2image library for PDF files"
                return result
            else:
           
                image = Image.open(file_path)
                text = pytesseract.image_to_string(
                    image, 
                    lang='+'.join(OCR_LANGUAGES)
                )
                result['text'] = clean_text(text)
                result['page_count'] = 1
        
        except Exception as e:
            self.logger.error(f"OCR extraction failed: {str(e)}")
            result['text'] = ''
        
        return result
